"""File processing utilities for different document types."""
from pathlib import Path
from typing import List, Dict, Any, Optional
import PyPDF2
from docx import Document


class FileProcessor:
    """Base class for file processing."""

    @staticmethod
    def get_processor(file_path: Path):
        """Get appropriate processor based on file extension."""
        extension = file_path.suffix.lower()

        if extension == '.pdf':
            return PDFProcessor()
        elif extension == '.txt':
            return TXTProcessor()
        elif extension == '.docx':
            return DOCXProcessor()
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    def extract_text(self, file_path: Path) -> str:
        """Extract text from file. Override in subclasses."""
        raise NotImplementedError

    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """
        Split text into overlapping chunks.

        Args:
            text: Text to chunk
            chunk_size: Target size of each chunk in characters
            overlap: Number of characters to overlap between chunks

        Returns:
            List of text chunks
        """
        if not text.strip():
            return []

        chunks = []
        words = text.split()

        current_chunk = []
        current_length = 0

        for word in words:
            word_length = len(word) + 1  # +1 for space

            if current_length + word_length > chunk_size and current_chunk:
                # Save current chunk
                chunk_text = " ".join(current_chunk)
                chunks.append(chunk_text)

                # Start new chunk with overlap
                # Keep last few words for context
                overlap_words = []
                overlap_length = 0
                for w in reversed(current_chunk):
                    if overlap_length + len(w) + 1 <= overlap:
                        overlap_words.insert(0, w)
                        overlap_length += len(w) + 1
                    else:
                        break

                current_chunk = overlap_words
                current_length = overlap_length

            current_chunk.append(word)
            current_length += word_length

        # Add remaining words
        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks


class PDFProcessor(FileProcessor):
    """Processor for PDF files."""

    def extract_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text = []

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text.append(page_text)

            return "\n\n".join(text)

        except Exception as e:
            raise ValueError(f"Error extracting PDF text: {e}")


class TXTProcessor(FileProcessor):
    """Processor for plain text files."""

    def extract_text(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue

            raise ValueError("Could not decode text file with any supported encoding")

        except Exception as e:
            raise ValueError(f"Error reading TXT file: {e}")


class DOCXProcessor(FileProcessor):
    """Processor for DOCX files."""

    def extract_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_text.append(" | ".join(row_text))

            # Combine all text
            all_text = paragraphs + table_text
            return "\n\n".join(all_text)

        except Exception as e:
            raise ValueError(f"Error extracting DOCX text: {e}")


def process_document(
    file_path: Path,
    chunk_size: int = 500,
    overlap: int = 50
) -> Dict[str, Any]:
    """
    Process a document file and return chunks with metadata.

    Args:
        file_path: Path to the document file
        chunk_size: Target size of each chunk
        overlap: Overlap between chunks

    Returns:
        Dictionary with chunks and metadata
    """
    processor = FileProcessor.get_processor(file_path)

    # Extract text
    text = processor.extract_text(file_path)

    if not text.strip():
        raise ValueError("No text content found in document")

    # Chunk text
    chunks = processor.chunk_text(text, chunk_size, overlap)

    return {
        "filename": file_path.name,
        "file_type": file_path.suffix.lower(),
        "file_size": file_path.stat().st_size,
        "text_length": len(text),
        "chunks": chunks,
        "num_chunks": len(chunks)
    }
